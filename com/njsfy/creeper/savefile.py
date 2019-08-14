#!/usr/bin/python
# coding: UTF-8
'''
Created on 2018年1月4日
存储文件

@author: shun
'''
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn


class SaveFile:
    '''
    classdocs
    '''


    def __init__(self, path):
        '''
        Constructor
        '''
        self.path = path
        self.document = Document()
        
        
    #添加一个新的章节    
    def addContent(self, content):
        
        font_name = u'宋体'
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(content)
        run.font.size = Pt(10)
        run.font.name = font_name
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), font_name)   
        return paragraph
        
    #往章节中增加内容
    def appendContent(self, paragraph, content):
        run = paragraph.add_run(content)
        font_name = u'宋体'
        run.font.size = Pt(10)
        run.font.name = font_name
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), font_name) 
        
    def addTitle(self, title):
        self.document.add_heading(title)
                 
    #保存文件    
    def saveDoc(self):
        self.document.save(self.path)
        